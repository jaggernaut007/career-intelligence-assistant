"""
Document Parser Service.

Handles parsing of PDF, DOCX, and plain text documents.
Uses puremagic library for file type detection (pure Python, no system dependencies).
"""

import io
import logging
from pathlib import Path
from typing import Optional, Union

import puremagic

from app.config import get_settings

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parses documents in various formats (PDF, DOCX, TXT)."""

    # Map MIME types to our internal types
    MIME_TO_TYPE = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/zip": "docx",  # DOCX files are ZIP archives
    }

    # Magic byte signatures for manual detection
    MAGIC_SIGNATURES = {
        b"%PDF": "pdf",
        b"PK\x03\x04": "docx",  # ZIP header (DOCX is a ZIP file)
    }

    def __init__(self):
        """Initialize document parser."""
        pass  # No initialization needed for puremagic

    def detect_file_type(self, file_input: Union[str, Path, io.BytesIO], filename: Optional[str] = None) -> str:
        """
        Detect file type using magic bytes.

        Args:
            file_input: Path to file or file-like object
            filename: Original filename (optional, for extension fallback)

        Returns:
            File type: 'pdf', 'docx', or 'txt'

        Raises:
            ValueError: If file type is unsupported
        """
        header = b""
        
        if isinstance(file_input, (str, Path)):
            file_path = Path(file_input)
            with open(file_path, "rb") as f:
                header = f.read(32)
            # Use puremagic helper for files
            try:
                results = puremagic.magic_file(str(file_path))
                if results and results[0].mime_type in self.MIME_TO_TYPE:
                    return self.MIME_TO_TYPE[results[0].mime_type]
            except Exception:
                pass
        else:
            # Stream
            pos = file_input.tell()
            header = file_input.read(32)
            file_input.seek(pos)
            
            # Use puremagic string helper
            try:
                # Read a bit more for puremagic - maybe 2KB
                pos = file_input.tell()
                start_bytes = file_input.read(2048)
                file_input.seek(pos)
                
                results = puremagic.magic_string(start_bytes)
                if results and results[0].mime_type in self.MIME_TO_TYPE:
                    return self.MIME_TO_TYPE[results[0].mime_type]
            except Exception:
                pass

        # Check our own magic signatures
        for signature, file_type in self.MAGIC_SIGNATURES.items():
            if header.startswith(signature):
                return file_type

        # Fallback to extension if provided
        if filename:
            ext = "." + filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
            if ext in (".txt", ".text"):
                return "txt"
        elif isinstance(file_input, (str, Path)):
            if Path(file_input).suffix.lower() in (".txt", ".text"):
                return "txt"

        # Try as plain text
        try:
            if isinstance(file_input, (str, Path)):
                with open(file_input, "rb") as f:
                    try:
                        f.read(1024).decode('utf-8')
                        return "txt"
                    except UnicodeDecodeError:
                        pass
            else:
                pos = file_input.tell()
                try:
                    file_input.read(1024).decode('utf-8')
                    file_input.seek(pos)
                    return "txt"
                except UnicodeDecodeError:
                    file_input.seek(pos)
        except Exception:
            pass

        raise ValueError("Unsupported file type")

    def validate_file_size(self, file_input: Union[str, Path, io.BytesIO]) -> None:
        """Validate file size."""
        settings = get_settings()
        size = 0
        if isinstance(file_input, (str, Path)):
            size = Path(file_input).stat().st_size
        else:
            current = file_input.tell()
            file_input.seek(0, io.SEEK_END)
            size = file_input.tell()
            file_input.seek(current)
            
        if size > settings.max_file_size_bytes:
            raise ValueError(f"File size exceeds limit")

    def validate_content_length(self, content: str) -> None:
        """
        Validate content length is under limit.

        Args:
            content: Text content

        Raises:
            ValueError: If content exceeds maximum length
        """
        settings = get_settings()
        if len(content) > settings.max_content_length:
            raise ValueError(
                f"Content length ({len(content)} chars) exceeds "
                f"maximum allowed length ({settings.max_content_length} chars)"
            )

    def parse(self, file_input: Union[str, Path, io.BytesIO], filename: Optional[str] = None) -> str:
        """
        Parse document and extract text content.

        Args:
            file_input: Path to document or file-like object
            filename: Optional filename for extension detection

        Returns:
            Extracted text content
        """
        # Validate file size
        self.validate_file_size(file_input)

        # Detect file type
        file_type = self.detect_file_type(file_input, filename)

        # Parse based on type
        content = ""
        if file_type == "pdf":
            content = self._parse_pdf(file_input)
        elif file_type == "docx":
            content = self._parse_docx(file_input)
        else:
            content = self._parse_text(file_input)

        if not content or not content.strip():
            raise ValueError("Document is empty or could not be parsed")

        self.validate_content_length(content)
        return content.strip()

    def _parse_pdf(self, file_input: Union[str, Path, io.BytesIO]) -> str:
        """Parse PDF document."""
        try:
            import fitz  # PyMuPDF

            if isinstance(file_input, (str, Path)):
                doc = fitz.open(str(file_input))
            else:
                # Stream
                doc = fitz.open(stream=file_input.read(), filetype="pdf")
                file_input.seek(0)

            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())

            doc.close()
            content = "\n".join(text_parts)

            if not content.strip():
                raise ValueError("PDF file is empty or contains only images")

            return content

        except Exception as e:
            if "empty" in str(e).lower():
                raise
            raise ValueError(f"Failed to parse PDF: {e}")

    def _parse_docx(self, file_input: Union[str, Path, io.BytesIO]) -> str:
        """Parse DOCX document."""
        try:
            from docx import Document

            if isinstance(file_input, (str, Path)):
                doc = Document(str(file_input))
            else:
                doc = Document(file_input)
                file_input.seek(0)

            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            content = "\n".join(text_parts)

            if not content.strip():
                raise ValueError("DOCX file is empty")

            return content

        except Exception as e:
            if "empty" in str(e).lower():
                raise
            raise ValueError(f"Failed to parse DOCX: {e}")

    def _parse_text(self, file_input: Union[str, Path, io.BytesIO]) -> str:
        """Parse plain text file."""
        try:
            content = ""
            if isinstance(file_input, (str, Path)):
                with open(file_input, "r", encoding="utf-8") as f:
                    content = f.read()
            else:
                content = file_input.read().decode('utf-8')
                file_input.seek(0)

            if not content.strip():
                raise ValueError("Text file is empty")

            return content

        except UnicodeDecodeError:
            # Try with different encoding
            try:
                if isinstance(file_input, (str, Path)):
                    with open(file_input, "r", encoding="latin-1") as f:
                        content = f.read()
                else:
                    file_input.seek(0)
                    content = file_input.read().decode('latin-1')
                    file_input.seek(0)
                
                if not content.strip():
                    raise ValueError("Text file is empty")
                return content
            except Exception:
                pass
            raise ValueError("Failed to decode text file")

    def parse_text(self, text: str) -> str:
        """
        Process and validate plain text input.

        Args:
            text: Raw text content

        Returns:
            Processed text

        Raises:
            ValueError: If text is empty
        """
        if not text or not text.strip():
            raise ValueError("Text content is empty")

        content = text.strip()
        self.validate_content_length(content)

        return content

    def parse_docx_content(self, content: bytes) -> str:
        """
        Parse DOCX content from bytes.

        Args:
            content: DOCX file content as bytes

        Returns:
            Extracted text

        Raises:
            ValueError: If content is empty or invalid
        """
        if not content:
            raise ValueError("DOCX content is empty")

        try:
            from docx import Document

            doc = Document(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            result = "\n".join(text_parts)

            if not result.strip():
                raise ValueError("DOCX content is empty")

            return result

        except Exception as e:
            if "empty" in str(e).lower():
                raise
            raise ValueError(f"Failed to parse DOCX content: {e}")


# Default parser instance
_default_parser: Optional[DocumentParser] = None


def get_document_parser() -> DocumentParser:
    """
    Get the default document parser instance.

    Returns:
        DocumentParser instance
    """
    global _default_parser
    if _default_parser is None:
        _default_parser = DocumentParser()
    return _default_parser
